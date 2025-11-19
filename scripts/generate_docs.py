from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def generate_gantt_docx(path: str = "docs/Gantt.docx"):
    doc = Document()
    doc.add_heading('Gantt Chart - Surface Attack Monitoring Tool', level=1)
    doc.add_paragraph('Timeline: 12 weeks (W1–W12). Adjust as needed.')
    doc.add_paragraph('Legend: ■ = planned work, blank = no work, [M] = milestone')

    # Header row: Task, Start, End, Dur, Depends, W01..W12
    cols = 5 + 12
    table = doc.add_table(rows=1, cols=cols)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    headers = ["Task", "Start", "End", "Dur", "Depends"] + [f"W{w:02d}" for w in range(1, 13)]
    for i, h in enumerate(headers):
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run.font.bold = True

    tasks = [
        {"name": "Initiation", "start": 1, "end": 2, "dur": 2, "dep": "-"},
        {"name": "Architecture & Design", "start": 2, "end": 4, "dur": 3, "dep": "Initiation"},
        {"name": "Backend Development", "start": 3, "end": 8, "dur": 6, "dep": "Arch & Design"},
        {"name": "Frontend & UI", "start": 4, "end": 7, "dur": 4, "dep": "Arch, Backend"},
        {"name": "Database & Persistence", "start": 3, "end": 5, "dur": 3, "dep": "Arch & Design"},
        {"name": "Security & Compliance", "start": 5, "end": 9, "dur": 5, "dep": "Backend, DB"},
        {"name": "Testing & QA", "start": 7, "end": 10, "dur": 4, "dep": "Backend, Frontend, Security"},
        {"name": "DevOps & Deployment", "start": 8, "end": 11, "dur": 4, "dep": "Backend, Testing"},
        {"name": "Documentation & Training", "start": 9, "end": 12, "dur": 4, "dep": "Testing, DevOps"},
        {"name": "Project Management", "start": 1, "end": 12, "dur": 12, "dep": "(ongoing)"},
    ]

    for t in tasks:
        cells = table.add_row().cells
        cells[0].text = t["name"]
        cells[1].text = f"W{t['start']}"
        cells[2].text = f"W{t['end']}"
        cells[3].text = str(t["dur"])
        cells[4].text = t["dep"]
        for w in range(1, 13):
            p = cells[4 + w].paragraphs[0]
            run = p.add_run("■" if t["start"] <= w <= t["end"] else "")
            run.font.name = 'Courier New'
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("\nNotes:")
    doc.add_paragraph("- Use Courier New in week columns to keep bars aligned.")
    doc.add_paragraph("- Replace W01–W12 with calendar dates as needed.")
    doc.save(path)


def generate_wbs_docx(path: str = "docs/WBS.docx"):
    doc = Document()
    doc.add_heading('Work Breakdown Structure (WBS) - Surface Attack Monitoring Tool', level=1)
    sections = [
        ("1. Initiation", [
            "1.1 Define scope and success criteria",
            "1.2 Identify stakeholders and roles",
            "1.3 Risk and dependency assessment",
            "1.4 Project schedule and milestones",
        ]),
        ("2. Architecture & Design", [
            "2.1 System architecture diagram (Flask + MongoDB + Email + Optional AI)",
            "2.2 Data model design (users, scans, results, logs)",
            "2.3 Security model (sessions, roles, admin)",
            "2.4 WBS and delivery plan approval",
        ]),
        ("3. Backend Development", [
            "3.1 Flask app structure and configuration",
            "3.2 Authentication (register/login, Google OAuth, optional 2FA)",
            "3.3 Admin panel (user management, cascade delete of scans)",
            "3.4 Scan ingestion (VirusTotal and optional ProjectDiscovery)",
            "3.5 Email service (SendGrid/SMTP for OTP and notices)",
            "3.6 Report generation (HTML + PDF via xhtml2pdf)",
            "3.7 API endpoints (scan submission, history, results)",
            "3.8 Error handling and logging",
        ]),
        ("4. Frontend & UI", [
            "4.1 Templates: index, dashboard, results, admin",
            "4.2 Navbar and responsive top-row actions",
            "4.3 Verify/OTP flow UX improvements",
            "4.4 Accessibility and responsive CSS (breakpoints)",
        ]),
        ("5. Database & Persistence", [
            "5.1 MongoDB connection, indexes (e.g., user_id, created_at)",
            "5.2 Data retention policy and cleanup jobs",
            "5.3 Seed data and admin bootstrap",
            "5.4 Backup/restore guidelines",
        ]),
        ("6. Security & Compliance", [
            "6.1 Input validation and sanitization",
            "6.2 Session security, cookies, CSRF",
            "6.3 Secrets management and environment loading",
            "6.4 Audit logging and admin actions",
        ]),
        ("7. Testing & QA", [
            "7.1 Unit tests (email send, cascade delete)",
            "7.2 Integration tests (auth, scans, history)",
            "7.3 Manual UAT checklist",
            "7.4 Performance and load checks",
        ]),
        ("8. DevOps & Deployment", [
            "8.1 Render setup (Procfile, build/start commands)",
            "8.2 Environment variables in Render",
            "8.3 Health checks and monitoring",
            "8.4 Release process and rollback plan",
        ]),
        ("9. Documentation & Training", [
            "9.1 README and setup guide",
            "9.2 Admin handbook",
            "9.3 User guide",
            "9.4 Operations runbook",
        ]),
        ("10. Project Management", [
            "10.1 Standups and status reporting",
            "10.2 Milestone tracking",
            "10.3 Risk review cadence",
            "10.4 Final delivery and retrospective",
        ]),
    ]

    for title, items in sections:
        doc.add_heading(title, level=2)
        for it in items:
            p = doc.add_paragraph(it)
            p.style = doc.styles['List Bullet']

    doc.save(path)


def generate_security_framework_docx(path: str = "docs/SecurityFramework.docx"):
    doc = Document()
    doc.add_heading('Security System Framework', level=1)
    doc.add_paragraph('Use this framework to structure security policies, controls, and operations. Customize per environment and regulatory needs.')

    sections = [
        ("1. Governance & Policy", [
            "Define security objectives, risk appetite, and scope",
            "Publish security policies: data classification, acceptable use, secure coding",
            "Establish roles and responsibilities (e.g., Admin, Security, DevOps)",
            "Maintain a risk register and control inventory",
        ]),
        ("2. Identity & Access Management (IAM)", [
            "Authentication: strong passwords, MFA, OAuth/OIDC where applicable",
            "Authorization: role-based access control (RBAC), least privilege",
            "Session management: secure cookies, timeouts, re-auth for critical actions",
            "Access reviews: periodic verification of admin and service accounts",
        ]),
        ("3. Data Protection", [
            "Encryption in transit (TLS) and at rest (database, backups)",
            "Key management: rotation, storage, access controls",
            "Data retention and deletion policies",
            "Secrets management: environment variables, secret stores; never commit .env",
        ]),
        ("4. Application Security", [
            "Input validation and output encoding (prevent XSS/Injection)",
            "Security headers (CSP, HSTS, X-Frame-Options)",
            "Dependency management: pin versions, scan for CVEs",
            "Secure logging and error handling (no sensitive data in logs)",
        ]),
        ("5. Network & Infrastructure Security", [
            "Segmentation and least-exposed services",
            "Firewalls/WAF and secure inbound/outbound rules",
            "TLS certificates lifecycle management",
            "System hardening (ports, services, configs)",
        ]),
        ("6. Endpoint & Device Security", [
            "Patch management and configuration baselines",
            "Anti-malware/EDR deployment and monitoring",
            "Restrict local admin rights and removable media",
        ]),
        ("7. Monitoring & Detection", [
            "Centralize logs (auth, admin actions, errors, scans)",
            "Alerting thresholds and on-call escalation",
            "Audit trails for admin and sensitive operations",
        ]),
        ("8. Incident Response", [
            "Playbooks for common events (credential compromise, malware, data leak)",
            "Triaging, containment, eradication, recovery steps",
            "Communication plans and legal/privacy notifications",
            "Post-incident review and corrective actions",
        ]),
        ("9. Vulnerability Management", [
            "Regular scanning, code review, and penetration tests",
            "Threat modeling for high-risk components",
            "Remediation SLAs by severity; track and verify fixes",
        ]),
        ("10. Secure Development Lifecycle (SDL)", [
            "Security requirements defined early",
            "Static analysis and secrets scanning in CI",
            "Peer reviews focused on security and privacy",
            "Release approvals and rollback procedures",
        ]),
        ("11. Compliance & Standards", [
            "Map controls to NIST CSF / ISO 27001 / local regulations",
            "Privacy requirements (data minimization, consent, user rights)",
            "Maintain evidence repository for audits",
        ]),
        ("12. Third-Party & Supply Chain", [
            "Vendor risk assessments and contractual security clauses",
            "SBOM and dependency provenance; pin and verify sources",
            "Monitor third-party service status and changes",
        ]),
        ("13. Cloud Security", [
            "Identity and access controls (IAM, keys, roles)",
            "Baseline configs, logging, and monitoring (e.g., cloud-native tools)",
            "Backups, DR, and secrets management in cloud services",
        ]),
        ("14. Physical Security", [
            "Access control to facilities and equipment",
            "Secure storage/disposal of media and backups",
        ]),
        ("15. Business Continuity & Disaster Recovery (BC/DR)", [
            "Define RTO/RPO targets and recovery strategies",
            "Regular backup testing and DR exercises",
            "Failover procedures and communication",
        ]),
        ("16. Training & Awareness", [
            "Role-based security training (developers, admins, users)",
            "Phishing and social engineering awareness",
            "Secure coding and infrastructure hardening training",
        ]),
        ("17. Metrics & Reporting", [
            "KPIs/KRIs (patching cadence, MTTR, incident counts)",
            "Security posture dashboards and risk trends",
            "Executive reporting cadence",
        ]),
    ]

    for title, items in sections:
        doc.add_heading(title, level=2)
        for it in items:
            p = doc.add_paragraph(it)
            p.style = doc.styles['List Bullet']

    # Add a simple NIST CSF mapping table
    doc.add_heading('Appendix: NIST CSF Mapping', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    table.cell(0, 0).text = 'Function'
    table.cell(0, 1).text = 'Example Controls'
    mappings = [
        ('Identify', 'Governance & Policy, Risk Register, Asset Inventory'),
        ('Protect', 'IAM, Data Protection, Application & Network Security, Training'),
        ('Detect', 'Monitoring & Detection, Audit Logs, Alerts'),
        ('Respond', 'Incident Response Playbooks, Communications, Triage'),
        ('Recover', 'BC/DR, Backups, Post-incident Improvements'),
    ]
    for i, (func, ctrls) in enumerate(mappings, start=1):
        table.cell(i, 0).text = func
        table.cell(i, 1).text = ctrls

    doc.add_paragraph('Notes: tailor controls to your stack and compliance obligations; ensure roles and SLAs are defined.')
    doc.save(path)

if __name__ == "__main__":
    generate_gantt_docx()
    generate_wbs_docx()
    generate_security_framework_docx()
    print("Generated: docs/Gantt.docx, docs/WBS.docx, docs/SecurityFramework.docx")